#!/usr/bin/env python3
"""
image_extractor.py
Module extraction images DOCX - VERSION CORRIGÉE
Conserve l'ordre d'apparition des images dans le document
"""

import os
from typing import Dict, Tuple, List
from pathlib import Path
from docx import Document
from PIL import Image
from io import BytesIO
import xml.etree.ElementTree as ET


def extract_all_images(docx_path: str, output_folder: str, base_name: str = "image") -> Tuple[Dict[int, str], Dict[int, dict]]:
    """
    Extrait TOUTES les images du DOCX dans l'ordre d'apparition
    VERSION CORRIGÉE - Conserve l'ordre exact des images
    
    Returns:
        tuple: (filenames_map, metadata_map)
        - filenames_map: {1: "image_001.png", 2: "image_002.jpg", ...}
        - metadata_map: {1: {"width": 800, "height": 600, "format": "PNG"}, ...}
    """
    os.makedirs(output_folder, exist_ok=True)
    
    doc = Document(docx_path)
    filenames_map = {}
    metadata_map = {}
    image_counter = 1
    
    # Créer un mapping des relations pour accès rapide
    rels_map = {}
    for rel_id, rel in doc.part.rels.items():
        if "image" in rel.target_ref:
            rels_map[rel_id] = rel
    
    # Parcourir le document dans l'ordre pour trouver les images
    for element in doc.element.body:
        if hasattr(element, 'xml'):
            # Chercher les images dans cet élément
            if './/pic:pic' in element.xml or 'graphic' in element.xml:
                # Parser le XML pour trouver les références d'images
                try:
                    # Chercher toutes les références d'images possibles
                    for match in ['r:embed="', 'r:link="', 'embed="', 'link="']:
                        if match in element.xml:
                            start_idx = element.xml.find(match)
                            while start_idx != -1:
                                start_idx += len(match)
                                end_idx = element.xml.find('"', start_idx)
                                if end_idx != -1:
                                    rel_id = element.xml[start_idx:end_idx]
                                    
                                    # Vérifier si c'est une relation d'image
                                    if rel_id in rels_map:
                                        rel = rels_map[rel_id]
                                        
                                        # Extraire l'image
                                        try:
                                            image_part = rel.target_part
                                            image_bytes = image_part.blob
                                            img = Image.open(BytesIO(image_bytes))
                                            
                                            # Déterminer l'extension
                                            ext = img.format.lower() if img.format else 'png'
                                            if ext == 'jpeg':
                                                ext = 'jpg'
                                            
                                            # Nom du fichier
                                            filename = f"{base_name}_{image_counter:03d}.{ext}"
                                            filepath = os.path.join(output_folder, filename)
                                            
                                            # Sauvegarder l'image
                                            with open(filepath, 'wb') as f:
                                                f.write(image_bytes)
                                            
                                            # Stocker le mapping
                                            filenames_map[image_counter] = filename
                                            metadata_map[image_counter] = {
                                                "width": img.width,
                                                "height": img.height,
                                                "format": img.format or "PNG",
                                                "rel_id": rel_id  # Garder le rel_id pour référence
                                            }
                                            
                                            # Marquer cette relation comme traitée
                                            del rels_map[rel_id]
                                            image_counter += 1
                                            
                                        except Exception as e:
                                            print(f"Erreur extraction image {rel_id}: {e}")
                                
                                # Chercher la prochaine occurrence
                                start_idx = element.xml.find(match, end_idx)
                                
                except Exception as e:
                    print(f"Erreur parsing XML: {e}")
                    continue
    
    return filenames_map, metadata_map


def get_image_positions(docx_path: str) -> Dict[int, int]:
    """
    Détermine la position de chaque image dans le document
    VERSION CORRIGÉE - Position exacte basée sur l'ordre d'apparition
    
    Returns:
        dict: {image_index: paragraph_position}
    """
    doc = Document(docx_path)
    positions = {}
    image_counter = 1
    para_position = 0
    
    # Parcourir tous les éléments du document
    for element in doc.element.body:
        # Vérifier si c'est un paragraphe
        from docx.oxml.text.paragraph import CT_P
        if isinstance(element, CT_P):
            from docx.text.paragraph import Paragraph
            para = Paragraph(element, doc)
            
            # Vérifier s'il y a une image dans ce paragraphe
            if para._element.xpath('.//pic:pic'):
                positions[image_counter] = para_position
                image_counter += 1
            
            para_position += 1
    
    return positions


def extract_images_with_positions(docx_path: str, output_folder: str) -> Tuple[Dict[int, str], Dict[int, dict], Dict[int, int]]:
    """
    Extrait les images avec leurs positions exactes dans le document
    
    Returns:
        tuple: (filenames_map, metadata_map, positions_map)
    """
    filenames_map, metadata_map = extract_all_images(docx_path, output_folder)
    positions_map = get_image_positions(docx_path)
    
    return filenames_map, metadata_map, positions_map


def create_image_mapping(docx_path: str) -> List[Tuple[int, str]]:
    """
    Crée un mapping ordonné des images dans le document
    
    Returns:
        list: [(position, image_ref_id), ...]
    """
    doc = Document(docx_path)
    image_mapping = []
    element_position = 0
    image_counter = 1
    
    for element in doc.element.body:
        from docx.oxml.text.paragraph import CT_P
        if isinstance(element, CT_P):
            from docx.text.paragraph import Paragraph
            para = Paragraph(element, doc)
            
            if para._element.xpath('.//pic:pic'):
                image_ref_id = f"__IMAGE_{image_counter}__"
                image_mapping.append((element_position, image_ref_id))
                image_counter += 1
        
        element_position += 1
    
    return image_mapping
